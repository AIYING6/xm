"""Build an anonymous 《航空学报》 Chinese submission draft from frozen sources.

The script deliberately does not accept author metadata.  It uses the official
Word template for page and style provenance, preserves the frozen figures, and
renders the manuscript as a double-blind, editable DOCX for final author review.
"""

from __future__ import annotations

import re
import sys
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "submission_packaging" / "official_source" / "航空学报_论文模板_20260226.docx"
SOURCE = ROOT / "paper_chinese" / "manuscript_zh.md"
BIB = ROOT / "paper_latex_3d_en" / "references.bib"
OUT = ROOT / "submission_packaging" / "航空学报_匿名投稿主文_预审版.docx"

FIGURE_DIR = ROOT / "paper_chinese" / "figures" / "publication"
FIGURES = {
    "图 1": FIGURE_DIR / "fig1_method_overview_publication.png",
    "图 2": FIGURE_DIR / "fig2_primary_recovery_publication.png",
}
EN_TITLES = {
    "图 1": "Fig. 1 Three-relation task graph for post-failure coordination",
    "图 2": "Fig. 2 Early task-chain recovery under matched failure exposure",
    "表 1": "Table 1 Primary results on the locked held-out evaluation",
    "表 2": "Table 2 Controlled component ablation",
    "表 3": "Table 3 Required boundary from zero-shot OOD evaluation",
}
EN_FIGURE_LEGENDS = {
    "图 1": "(a) Scout–Relay–Attack–Target heterogeneous task scenario and the three relations. (b) Three-relation task graphs before and after relay failure; the failure makes the corresponding relations unavailable, and the attack window is not a fourth relation. (c) Local observations and available graphs are processed by relation-specific edge-aware attention, Gate Prior, static Role-Pair modulation and union-graph residual fusion before decentralized actor execution. Blue dotted, green dashed and orange dash-dotted lines denote perception, environment-delivered communication and task support, respectively. Relation adjacency is an aggregation mask, not a learned physical communication switch.",
    "图 2": "(a) Full Kaplan–Meier recovery curves for EA-RG, MAPPO, HAPPO and the wider single-graph baseline; the fine dashed line marks the end of the pre-specified active node-failure window at 80 steps. (b) Detail from 0 to 35 steps, where the primary separation occurs. (c) Seed-level RMST80 differences between EA-RG and MAPPO and the pooled hierarchical paired-bootstrap 95% interval; negative values indicate earlier recovery by EA-RG. Curves summarize three independent training seeds and 600 failure-exposed episodes per method; unrecovered episodes are right-censored.",
}
EN_TABLE_LEGENDS = {
    "表 1": "Recovery and Success are means ± sample standard deviations over three independent training seeds. Conditional mean recovery time is defined only for failure-exposed episodes that recovered and cannot replace RMST with right censoring. RMST80 is the pre-specified P1 comparison between EA-RG and MAPPO; other RMST80 values are reported in the Supplementary Information.",
    "表 2": "Means ± sample standard deviations are based on three independent training seeds. Conditional time and RMST target different estimands; mechanism traces and seed-level details are reported in the Supplementary Information.",
    "表 3": "Values are locked family-level summaries used only to describe shift dependence; they are not p values or evidence of universal generalization.",
}
EN_ABSTRACT = (
    "Critical relay failures can interrupt task chains in heterogeneous UAV formations operating with intermittent sensing and constrained communication. "
    "Terminal success alone does not reveal when usable coordination returns while the failure persists. We study post-failure task-chain recovery and propose Edge-Aware Role-Graph MAPPO (EA-RG). "
    "Under centralized training and decentralized execution, EA-RG represents perception, environment-delivered communication and task support as three relations, which are aggregated by relation-specific edge-aware graph attention and a union-graph residual path. "
    "In a locked nominal held-out evaluation with matched failure-exposed episodes, we used Kaplan–Meier curves and restricted mean survival time (RMST). Across three independent training seeds, EA-RG had an RMST of 11.81 steps versus 15.51 steps for MAPPO in the pre-specified 80-step node-failure window; the hierarchical paired-bootstrap 95% interval for their contrast was [−7.16, −1.05] steps. "
    "At the common 220-step follow-up horizon, EA-RG remained lower than MAPPO, whereas comparisons with HAPPO and a wider single-graph baseline were not directionally consistent. Component ablations provided limited support, and zero-shot shifts were shift-family dependent. "
    "Thus, the contribution is restricted to earlier task-chain recovery under the locked nominal condition, not universal baseline or out-of-distribution superiority."
)
EN_KEYWORDS = "heterogeneous unmanned aerial vehicles; multi-agent reinforcement learning; post-failure recovery; task graph; multi-relational graph attention; constrained communication; survival analysis"


def set_columns(section, count: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.first_child_found_in("w:cols")
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), "425")


def set_run_font(run, size: float, bold: bool = False, english: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman" if english else "宋体"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")


def add_paragraph(doc, text: str = "", *, size: float = 9, bold: bool = False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, english: bool = False,
                  first_indent: bool = False, space_after: float = 1.5):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.15
    if first_indent:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size, bold, english)
    return p


def clean_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\\\((.*?)\\\)", r"\1", text)
    replacements = {
        r"\tau": "τ", r"\pm": "±", r"\times": "×", r"\le": "≤", r"\ge": "≥",
        r"\mathrm": "", r"\text": "", r"\_": "_", "{": "", "}": "",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return re.sub(r"\s+", " ", text).strip()


def citation_order(markdown: str) -> OrderedDict[str, int]:
    ordered: OrderedDict[str, int] = OrderedDict()
    for group in re.findall(r"\[@([^\]]+)\]", markdown):
        # The opening @ is consumed by the outer expression; restore it so
        # the first key and later semicolon-separated keys are treated alike.
        for key in re.findall(r"@([A-Za-z0-9_:-]+)", "@" + group):
            ordered.setdefault(key, len(ordered) + 1)
    return ordered


def citation_label(keys: list[str], order: OrderedDict[str, int]) -> str:
    nums = sorted(order[key] for key in keys)
    ranges, begin, end = [], nums[0], nums[0]
    for num in nums[1:]:
        if num == end + 1:
            end = num
        else:
            ranges.append(str(begin) if begin == end else f"{begin}-{end}")
            begin = end = num
    ranges.append(str(begin) if begin == end else f"{begin}-{end}")
    return "[" + ",".join(ranges) + "]"


def replace_citations(text: str, order: OrderedDict[str, int]) -> str:
    def replace(match: re.Match[str]) -> str:
        keys = re.findall(r"@([A-Za-z0-9_:-]+)", "@" + match.group(1))
        return citation_label(keys, order)
    return re.sub(r"\[@([^\]]+)\]", replace, text)


def parse_bib(path: Path) -> dict[str, dict[str, str]]:
    raw = path.read_text(encoding="utf-8")
    entries: dict[str, dict[str, str]] = {}
    for match in re.finditer(r"@(\w+)\s*\{\s*([^,]+),([\s\S]*?)\n\}", raw):
        entry_type, key, body = match.groups()
        fields = {"type": entry_type.lower()}
        for fmatch in re.finditer(r"(\w+)\s*=\s*[\{\"](.*?)[\}\"],?\s*(?=\w+\s*=|$)", body, flags=re.S):
            field, value = fmatch.groups()
            fields[field.lower()] = re.sub(r"\s+", " ", value).strip()
        entries[key.strip()] = fields
    return entries


def bib_text(field: str) -> str:
    field = clean_markdown(field)
    field = field.replace("{", "").replace("}", "")
    return field.replace("\\&", "&")


def format_reference(index: int, entry: dict[str, str]) -> str:
    authors = bib_text(entry.get("author", ""))
    authors = authors.replace(" and ", ", ")
    title = bib_text(entry.get("title", ""))
    year = bib_text(entry.get("year", ""))
    typ = entry.get("type", "article")
    if typ in {"article", "inproceedings", "conference"}:
        medium = "J" if typ == "article" else "C"
        venue = bib_text(entry.get("journal", entry.get("booktitle", "")))
        volume = bib_text(entry.get("volume", ""))
        number = bib_text(entry.get("number", ""))
        pages = bib_text(entry.get("pages", "")).replace("--", "-")
        tail = f"{venue}, {year}"
        if volume:
            tail += f", {volume}"
            if number:
                tail += f"({number})"
        if pages:
            tail += f": {pages}"
        text = f"[{index}] {authors}. {title}[{medium}]. {tail}."
    else:
        publisher = bib_text(entry.get("publisher", ""))
        text = f"[{index}] {authors}. {title}[M]. {publisher}, {year}."
    doi = bib_text(entry.get("doi", ""))
    if doi:
        text += f" DOI: {doi}."
    return text


def add_wide_section(doc, columns: int):
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(section, columns)
    return section


def add_table(doc, rows: list[list[str]]) -> None:
    add_wide_section(doc, 1)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(clean_markdown(value))
            set_run_font(run, 7.2, bold=(r == 0))
    add_wide_section(doc, 2)


def add_figure(doc, label: str) -> None:
    path = FIGURES[label]
    if not path.exists():
        raise FileNotFoundError(path)
    add_wide_section(doc, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(17.3))
    add_wide_section(doc, 2)


def add_caption(doc, label: str, chinese: str) -> None:
    add_paragraph(doc, chinese, size=8, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1)
    if label.startswith("图"):
        add_paragraph(doc, EN_TITLES[label] + ". " + EN_FIGURE_LEGENDS[label], size=8, english=True, space_after=2)
    else:
        add_paragraph(doc, EN_TITLES[label] + ". " + EN_TABLE_LEGENDS[label], size=8, english=True, space_after=2)


def add_front_matter(doc, title: str, abstract: str, keywords: str) -> None:
    p = add_paragraph(doc, title, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    for run in p.runs:
        set_run_font(run, 18, True)
    add_paragraph(doc, "摘要：" + abstract, size=9, first_indent=False, space_after=3)
    add_paragraph(doc, "关键词：" + keywords, size=9, space_after=8)
    p = add_paragraph(doc, "Multi-relational graph decision-making for post-failure task-chain recovery in heterogeneous unmanned aerial vehicles", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, english=True, space_after=5)
    for run in p.runs:
        set_run_font(run, 14, True, True)
    add_paragraph(doc, "Abstract: " + EN_ABSTRACT, size=9, english=True, space_after=3)
    add_paragraph(doc, "Key words: " + EN_KEYWORDS, size=9, english=True, space_after=8)


def build() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    order = citation_order(markdown)
    bib = parse_bib(BIB)
    missing = [key for key in order if key not in bib]
    if missing:
        raise RuntimeError(f"Cited keys missing from BibTeX: {missing}")

    doc = Document(TEMPLATE)
    body = doc._element.body
    for element in list(body):
        if element.tag != qn("w:sectPr"):
            body.remove(element)
    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.title = ""
    props.subject = ""
    props.comments = ""
    props.keywords = ""
    set_columns(doc.sections[0], 1)

    title = re.search(r"^#\s+(.+)$", markdown, flags=re.M).group(1)
    abstract = re.search(r"## 摘要\s*\n\n(.+?)\n\n\*\*关键词：\*\*\s*([^\n]+)", markdown, flags=re.S)
    if not abstract:
        raise RuntimeError("Cannot locate Chinese abstract and keywords")
    add_front_matter(doc, title, clean_markdown(abstract.group(1)), clean_markdown(abstract.group(2)))

    doc.add_section(WD_SECTION.NEW_PAGE)
    set_columns(doc.sections[-1], 2)
    # The title, abstract and keywords have already been rendered above.  Only
    # convert the scientific body, starting at the first numbered section.
    body_match = re.search(r"(?m)^## 1 引言\s*$", markdown)
    if body_match is None:
        raise RuntimeError("Cannot locate the first manuscript section")
    lines = markdown[body_match.start():].splitlines()
    i = 0
    figure_pending = None
    table_number = 0
    while i < len(lines):
        line = lines[i].strip()
        i += 1
        if not line or line == "---" or line.startswith(("title:", "language:", "status:", "evidence_freeze:")):
            continue
        if line.startswith("# ") or line == "## 摘要" or line.startswith("**关键词：") or line.startswith("> "):
            continue
        if line.startswith("![") and "图" in line:
            figure_pending = re.search(r"图\s+\d+", line).group(0)
            add_figure(doc, figure_pending)
            continue
        if line.startswith("**图 ") or line.startswith("**表 "):
            label = re.search(r"(?:图|表)\s+\d+", line).group(0)
            add_caption(doc, label, clean_markdown(line))
            figure_pending = None
            continue
        if line.startswith("## "):
            add_paragraph(doc, clean_markdown(line[3:]), size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3)
            continue
        if line.startswith("### "):
            add_paragraph(doc, clean_markdown(line[4:]), size=10.5, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
            continue
        if line.startswith("|"):
            rows = [line]
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            parsed = [[cell.strip() for cell in row.strip("|").split("|")] for row in rows]
            parsed = [row for idx, row in enumerate(parsed) if idx != 1]
            add_table(doc, parsed)
            table_number += 1
            continue
        if line.startswith("- "):
            add_paragraph(doc, "• " + clean_markdown(replace_citations(line[2:], order)), size=9, first_indent=False, space_after=1)
            continue
        text = clean_markdown(replace_citations(line, order))
        add_paragraph(doc, text, size=9, first_indent=True, space_after=1.5)

    add_paragraph(doc, "参考文献", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3)
    for key, index in order.items():
        add_paragraph(doc, format_reference(index, bib[key]), size=8, english=True, first_indent=False, space_after=0.5)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Citations: {len(order)}; figures: {len(FIGURES)}; tables: {table_number}")


if __name__ == "__main__":
    try:
        build()
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise
